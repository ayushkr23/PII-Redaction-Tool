import argparse
import os
from docx import Document
from detectors import PIIDetector
from replacer import PIIReplacer

from collections import Counter

class PIIRedactor:
    def __init__(self):
        self.detector = PIIDetector()
        self.replacer = PIIReplacer()
        self.stats = Counter()

    def redact_text(self, text: str) -> str:
        """
        Detects PII in a string and replaces it with fake data.
        """
        if not text.strip():
            return text
            
        entities = self.detector.detect(text)
        if not entities:
            return text
            
        # Update stats
        for entity in entities:
            self.stats[entity['entity_type']] += 1
            
        # Replace from end to start to avoid index shifting
        redacted_text = text
        for entity in sorted(entities, key=lambda x: x['start'], reverse=True):
            start = entity['start']
            end = entity['end']
            original_val = text[start:end]
            fake_val = self.replacer.get_replacement(entity['entity_type'], original_val)
            redacted_text = redacted_text[:start] + fake_val + redacted_text[end:]
            
        return redacted_text

    def redact_document(self, input_path: str, output_path: str):
        """
        Reads a docx, redacts PII, and saves it.
        """
        doc = Document(input_path)
        
        # Redact paragraphs
        for para in doc.paragraphs:
            self._redact_paragraph(para)
            
        # Redact tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._redact_paragraph(para)
                        
        doc.save(output_path)
        print(f"Redacted document saved to: {output_path}")

    def _redact_paragraph(self, paragraph):
        """
        Redacts PII in a paragraph while trying to maintain formatting.
        Note: If PII spans across multiple runs, a simple run-by-run replacement might fail.
        To ensure high detection recall, we detect on the full paragraph text.
        If PII is found, we do a simplistic text replacement across runs.
        """
        full_text = paragraph.text
        if not full_text.strip():
            return
            
        redacted_full_text = self.redact_text(full_text)
        
        if redacted_full_text != full_text:
            # PII was found. 
            # Replacing text in python-docx while perfectly preserving runs is complex.
            # We clear all runs and put the redacted text in the first run to preserve basic paragraph formatting.
            # This is a known tradeoff documented in the README.
            if len(paragraph.runs) > 0:
                style = paragraph.runs[0].style
                for i in range(len(paragraph.runs)):
                    paragraph.runs[i].text = ""
                paragraph.runs[0].text = redacted_full_text
            else:
                paragraph.text = redacted_full_text


def main():
    parser = argparse.ArgumentParser(description="PII Redaction Tool")
    parser.add_argument("--input", type=str, required=True, help="Path to input DOCX file")
    parser.add_argument("--output", type=str, required=True, help="Path to output redacted DOCX file")
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: Input file {args.input} not found.")
        return
        
    print(f"Starting redaction on {args.input}...")
    redactor = PIIRedactor()
    redactor.redact_document(args.input, args.output)
    print("Redaction complete.")

if __name__ == "__main__":
    main()
